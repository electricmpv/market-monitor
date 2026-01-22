#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
🔍 机会猎手 v2.0 - 融资、创业、技术突破监控
Author: 电动面包
Purpose: 发现融资项目、创业团队、技术突破
"""

import os
import datetime
import time
import hashlib
import json
import sys
from pathlib import Path

try:
    from google import genai
    import requests
    import chromadb
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError as e:
    print(f"❌ 依赖库缺失: {e}")
    sys.exit(1)

# ==================== 🛠️ 用户配置区 ====================

GEMINI_KEY = os.getenv('GEMINI_API_KEY', '填入自己的API-key')
PUSHPLUS_TOKEN = os.getenv('PUSHPLUS_TOKEN', '填入自己的TOKEN')
GITHUB_TOKEN = os.getenv('GITHUB_TOKEN', '')

YOUR_PORT = int(os.getenv('PROXY_PORT', 19828))
USE_PROXY = YOUR_PORT > 0

# GitHub搜索关键词 - 精确版
GITHUB_KEYWORDS = [
    # AI Agent框架
    'AI agent framework', 'LLM agent', 'autonomous agent',
    'agent orchestration', 'multi-agent',
    
    # RAG系统
    'RAG pipeline', 'retrieval augmented', 'vector database',
    'semantic search', 'knowledge graph',
    
    # 提示工程
    'prompt engineering', 'prompt optimization',
    'prompt template', 'few-shot learning',
    
    # 自动化
    'workflow automation', 'task automation',
    'browser automation', 'API automation',
    
    # 新工具
    'LLM framework', 'AI framework', 'generative AI',
    'DeepSeek', 'Claude integration', 'GPT wrapper'
]

# GitHub过滤
MIN_STARS = 300
DAYS_SINCE_UPDATE = 90

# Hacker News关键词
HN_KEYWORDS = [
    'AI', 'machine learning', 'LLM', 'GPT', 'Claude',
    'startup', 'funding', 'Series A', 'Series B',
    'open source', 'breakthrough', 'SOTA'
]

# =======================================================================

if USE_PROXY:
    PROXY_URL = f'http://127.0.0.1:{YOUR_PORT}'
    os.environ['http_proxy'] = PROXY_URL
    os.environ['https_proxy'] = PROXY_URL
    print(f"📡 代理已启用: {PROXY_URL}")

DATA_DIR = Path('./my_market_brain')
DATA_DIR.mkdir(exist_ok=True)

print(f"🔍 机会猎手 v2.0 启动... [时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}]")

try:
    gemini_client = genai.Client(api_key=GEMINI_KEY)
    chroma_client = chromadb.PersistentClient(path=str(DATA_DIR))
    opportunity_collection = chroma_client.get_or_create_collection(name="opportunities_v2")
    print("✅ 所有组件加载完毕")
except Exception as e:
    print(f"❌ 初始化失败: {e}")
    sys.exit(1)

current_session_opportunities = []

# ==================== 工具函数 ====================

def create_robust_session():
    """创建抗网络波动的会话"""
    session = requests.Session()
    if USE_PROXY:
        session.proxies = {"http": PROXY_URL, "https": PROXY_URL}
    return session

http = create_robust_session()

def save_opportunity(source, title, description, link, metadata):
    """保存机会"""
    try:
        content = f"{source}: {title} | {description}"
        content_fingerprint = hashlib.md5(content.encode('utf-8')).hexdigest()
        doc_id = f"OPP_{source}_{content_fingerprint}"
        
        # 检查重复
        existing = opportunity_collection.get(ids=[doc_id])
        if existing and existing['ids']:
            return False
        
        current_time = datetime.datetime.now().isoformat()
        opportunity_collection.upsert(
            documents=[content],
            metadatas=[{
                "source": source,
                "title": title,
                "type": metadata.get('type', 'unknown'),
                "time": current_time,
                "link": link
            }],
            ids=[doc_id]
        )
        
        current_session_opportunities.append({
            'source': source,
            'title': title,
            'description': description,
            'link': link,
            'metadata': metadata,
            'time': current_time
        })
        
        print(f"  💡 [{source}] {title[:50]}...")
        return True
    except Exception as e:
        print(f"  ⚠️ 保存失败: {e}")
        return False

def hunt_github():
    """GitHub项目猎手"""
    print("\n🐙 [1/2] 正在扫描 GitHub...")
    count = 0
    
    headers = {
        'Accept': 'application/vnd.github.v3+json',
        'User-Agent': 'MarketHunter/v2'
    }
    if GITHUB_TOKEN:
        headers['Authorization'] = f'token {GITHUB_TOKEN}'
    
    try:
        for keyword in GITHUB_KEYWORDS[:5]:  # 每次选5个关键词
            print(f"  🔍 搜索: {keyword}")
            
            api_url = f"https://api.github.com/search/repositories?q={keyword}+stars:>{MIN_STARS}&sort=updated&order=desc&per_page=3"
            
            try:
                resp = http.get(api_url, headers=headers, timeout=15)
                
                if resp.status_code == 200:
                    items = resp.json().get('items', [])
                    
                    for item in items:
                        updated_at = item['updated_at'][:10]
                        last_update = datetime.datetime.strptime(updated_at, "%Y-%m-%d")
                        days_diff = (datetime.datetime.now() - last_update).days
                        
                        if days_diff > DAYS_SINCE_UPDATE:
                            continue
                        
                        if save_opportunity(
                            source="GitHub",
                            title=item['full_name'],
                            description=item['description'] or "No description",
                            link=item['html_url'],
                            metadata={
                                'type': 'OpenSource',
                                'stars': item['stargazers_count'],
                                'language': item['language'],
                                'updated': updated_at
                            }
                        ):
                            count += 1
                
                elif resp.status_code == 403:
                    print("⛔ GitHub API 频率超限")
                    break
                
                time.sleep(1)
                
            except Exception as e:
                print(f"     ⚠️ 搜索出错: {e}")
                continue
    
    except Exception as e:
        print(f"❌ GitHub 扫描失败: {e}")
    
    return count

def hunt_hacker_news():
    """Hacker News机会猎手"""
    print("\n📰 [2/2] 正在扫描 Hacker News...")
    count = 0
    
    try:
        resp = requests.get(
            'https://hacker-news.firebaseio.com/v0/topstories.json',
            timeout=10
        )
        top_ids = resp.json()[:15]
        
        for item_id in top_ids:
            try:
                item = requests.get(
                    f'https://hacker-news.firebaseio.com/v0/item/{item_id}.json',
                    timeout=10
                ).json()
                
                if item and item.get('score', 0) >= 150:
                    title = item.get('title', '')
                    text = item.get('text', '')
                    url = item.get('url', '')
                    
                    # 检查是否包含机会关键词
                    content_lower = (title + text).lower()
                    is_opportunity = False
                    opp_type = 'News'
                    
                    if any(kw in content_lower for kw in ['funding', 'series', 'raised', 'investment']):
                        is_opportunity = True
                        opp_type = 'Funding'
                    elif any(kw in content_lower for kw in ['startup', 'founded', 'launch']):
                        is_opportunity = True
                        opp_type = 'Startup'
                    elif any(kw in content_lower for kw in ['breakthrough', 'SOTA', 'new', 'release']):
                        is_opportunity = True
                        opp_type = 'Technology'
                    
                    if is_opportunity:
                        if save_opportunity(
                            source="HackerNews",
                            title=title,
                            description=text[:200],
                            link=url,
                            metadata={
                                'type': opp_type,
                                'score': item.get('score', 0)
                            }
                        ):
                            count += 1
                
                time.sleep(0.5)
            except:
                pass
    
    except Exception as e:
        print(f"❌ HN 扫描失败: {e}")
    
    return count

def analyze_opportunities_ai(raw_data):
    """AI分析机会"""
    print("\n🧠 正在用AI分析机会...")
    
    prompt = f"""
# Role: Investment & Startup Analyst
# Task: 从技术新闻和开源项目中识别商业机会

## 分析维度
1. **融资信号** - 哪些创业公司获得融资？为什么？
2. **技术趋势** - 哪些技术方向在升温？
3. **工具机会** - 哪些开源项目可能商业化？
4. **市场缺口** - 还有哪些未被满足的需求？

## 数据
{raw_data}

## 输出格式
### 🎯 Top 5 机会

1. [机会名称]
   - 类型: [融资/技术/工具/市场]
   - 核心价值: ...
   - 为什么重要: ...
   - 你的行动: ...

2. [机会名称]
   ...

### 📊 趋势总结
- 最热话题: ...
- 融资热度: ...
- 技术方向: ...
"""
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = gemini_client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt
            )
            return response.text
        except Exception as e:
            print(f"⚠️ 分析尝试 {attempt+1} 失败: {e}")
            if attempt < max_retries - 1:
                time.sleep(5)
    
    return "❌ AI分析失败"

def deliver_report(content):
    """交付报告"""
    if content.startswith("❌"):
        print(f"\n🚫 {content}")
        return
    
    today = datetime.date.today().strftime("%Y-%m-%d")
    filename = f"Opportunities_Report_{today}.docx"
    
    # 保存Word
    try:
        doc = Document()
        doc.add_heading(f'🔍 机会发现报告 - {today}', 0)
        doc.add_paragraph(f"生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"发现机会数: {len(current_session_opportunities)}")
        doc.add_paragraph("=" * 50)
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            else:
                doc.add_paragraph(line)
        
        doc.save(filename)
        print(f"\n💾 ✅ 报告已生成: {filename}")
    except Exception as e:
        print(f"❌ Word生成失败: {e}")
    
    # 推送微信
    if PUSHPLUS_TOKEN and PUSHPLUS_TOKEN != '填入自己的TOKEN':
        try:
            print("📨 正在推送到微信...")
            wechat_body = f"# 🔍 机会发现报告 ({today})\n\n{content}"
            
            requests.post(
                'http://www.pushplus.plus/send',
                json={
                    "token": PUSHPLUS_TOKEN,
                    "title": f"【机会】{today}",
                    "content": wechat_body,
                    "template": "markdown"
                },
                timeout=10
            )
            print("📨 ✅ 微信推送完成")
        except Exception as e:
            print(f"⚠️ 推送失败: {e}")

# ==================== 主程序 ====================

def main():
    print("\n" + "="*60)
    print("🚀 开始机会猎手循环")
    print("="*60)
    
    c1 = hunt_github()
    c2 = hunt_hacker_news()
    
    total = c1 + c2
    print(f"\n📊 本次发现机会数: {total}")
    
    if total > 0:
        raw_opps = "\n".join([
            f"【{o['source']}】{o['title']}: {o['description']}"
            for o in current_session_opportunities
        ])
        
        analysis = analyze_opportunities_ai(raw_opps)
        deliver_report(analysis)
    else:
        print("🤷 未发现新机会")
    
    print("\n✅ 机会猎手循环完成")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n⚠️ 用户中断")
    except Exception as e:
        print(f"\n❌ 程序错误: {e}")
        import traceback
        traceback.print_exc()
