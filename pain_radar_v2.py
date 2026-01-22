#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ğŸ¯ ç—›ç‚¹é›·è¾¾ v2.0 - ç²¾ç¡®å¸‚åœºæœºä¼šç›‘æ§ç³»ç»Ÿ
Author: ç”µåŠ¨é¢åŒ… (AI Solopreneur)
Purpose: ä»å…¨ç½‘ç”¨æˆ·åæ§½ä¸­æå–å¯å•†ä¸šåŒ–çš„å¸‚åœºæœºä¼š
"""

import os
import asyncio
import datetime
import time
import hashlib
import random
import json
import sys
from pathlib import Path

try:
    from google import genai
    from twikit import Client
    import requests
    import chromadb
    from docx import Document
except ImportError as e:
    print(f"âŒ ä¾èµ–åº“ç¼ºå¤±: {e}")
    print("è¯·è¿è¡Œ: pip install google-genai twikit requests chromadb python-docx")
    sys.exit(1)

# ==================== ğŸ› ï¸ ç”¨æˆ·é…ç½®åŒº ====================

# APIå¯†é’¥é…ç½®
GEMINI_KEY = os.getenv('GEMINI_API_KEY', 'å¡«å…¥è‡ªå·±çš„API-key')
PUSHPLUS_TOKEN = os.getenv('PUSHPLUS_TOKEN', 'å¡«å…¥è‡ªå·±çš„TOKEN')
GITHUB_TOKEN = os.getenv('GITHUB_TOKEN', '')

# ç½‘ç»œé…ç½®
YOUR_PORT = int(os.getenv('PROXY_PORT', 19828))  # æ¢¯å­ç«¯å£ï¼Œå¦‚æœä¸éœ€è¦ä»£ç†è®¾ä¸º0
USE_PROXY = YOUR_PORT > 0

# ç›‘æ§é…ç½® - ç²¾ç¡®å…³é”®è¯ï¼ˆå·²ä¼˜åŒ–ï¼‰
PAIN_KEYWORDS = {
    'ChatGPT': [
        'can\'t', 'doesn\'t work', 'error', 'failed',
        'slow', 'expensive', 'confusing', 'limitation'
    ],
    'Claude': [
        'can\'t', 'doesn\'t support', 'bug', 'api down',
        'rate limit', 'context window', 'expensive'
    ],
    'DeepSeek': [
        'slow', 'error', 'hallucination', 'can\'t',
        'doesn\'t work', 'quality issue'
    ],
    'Cursor': [
        'bug', 'crash', 'doesn\'t work', 'slow',
        'indexing fail', 'completion wrong'
    ],
    'Midjourney': [
        'hands weird', 'broken', 'ugly', 'consistency',
        'text fail', 'quality drop'
    ],
    'Sora': [
        'physics fail', 'movement unnatural', 'face melting',
        'flicker', 'artifact', 'quality'
    ]
}

# åƒåœ¾è¯é»‘åå•
SPAM_FILTERS = [
    '100+ AI Tools', 'Check my bio', 'Sign up now',
    'Top 10 tools', 'Affiliate', 'Giveaway', 'NFT',
    'crypto', 'bitcoin', 'follow me', 'DM me'
]

# =======================================================================

# ç¯å¢ƒé…ç½®
if USE_PROXY:
    PROXY_URL = f'http://127.0.0.1:{YOUR_PORT}'
    os.environ['http_proxy'] = PROXY_URL
    os.environ['https_proxy'] = PROXY_URL
    os.environ['HTTP_PROXY'] = PROXY_URL
    os.environ['HTTPS_PROXY'] = PROXY_URL
    print(f"ğŸ“¡ ä»£ç†å·²å¯ç”¨: {PROXY_URL}")
else:
    print("ğŸ“¡ ä»£ç†å·²ç¦ç”¨ï¼ˆç›´è¿æ¨¡å¼ï¼‰")

# åˆ›å»ºæ•°æ®ç›®å½•
DATA_DIR = Path('./my_market_brain')
DATA_DIR.mkdir(exist_ok=True)

print(f"ğŸ¯ ç—›ç‚¹é›·è¾¾ v2.0 å¯åŠ¨... [æ—¶é—´: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}]")

# åˆå§‹åŒ–ç»„ä»¶
try:
    gemini_client = genai.Client(api_key=GEMINI_KEY)
    chroma_client = chromadb.PersistentClient(path=str(DATA_DIR))
    pain_collection = chroma_client.get_or_create_collection(name="pain_points_v2")
    print("âœ… æ‰€æœ‰ç»„ä»¶åŠ è½½å®Œæ¯•")
except Exception as e:
    print(f"âŒ åˆå§‹åŒ–å¤±è´¥: {e}")
    sys.exit(1)

current_session_pains = []

# ==================== å·¥å…·å‡½æ•° ====================

def is_spam(text):
    """åƒåœ¾å†…å®¹æ£€æµ‹"""
    text_lower = text.lower()
    for spam_word in SPAM_FILTERS:
        if spam_word.lower() in text_lower:
            return True
    return False

def save_pain(source, author, content, product):
    """ä¿å­˜ç—›ç‚¹åˆ°æ•°æ®åº“"""
    try:
        if is_spam(content):
            return False
            
        current_time = datetime.datetime.now().isoformat()
        content_fingerprint = hashlib.md5(content.encode('utf-8')).hexdigest()
        doc_id = f"PAIN_{source}_{product}_{content_fingerprint}"
        
        # æ£€æŸ¥æ˜¯å¦å·²å­˜åœ¨
        existing = pain_collection.get(ids=[doc_id])
        if existing and existing['ids']:
            return False
        
        pain_collection.upsert(
            documents=[content],
            metadatas=[{
                "source": source,
                "author": str(author),
                "product": product,
                "type": "pain",
                "time": current_time
            }],
            ids=[doc_id]
        )
        current_session_pains.append({
            'source': source,
            'author': author,
            'product': product,
            'content': content,
            'time': current_time
        })
        print(f"  ğŸ©¸ [{product}] {content[:50]}...")
        return True
    except Exception as e:
        print(f"  âš ï¸ ä¿å­˜å¤±è´¥: {e}")
        return False

async def scan_twitter():
    """æ‰«æTwitterç—›ç‚¹"""
    print("\nğŸ¦ [1/3] æ­£åœ¨æ‰«æ Twitter...")
    client = Client(language='en-US')
    count = 0
    
    try:
        client.load_cookies('cookies.json')
        
        # æ„å»ºæœç´¢æŸ¥è¯¢
        search_queries = []
        for product, keywords in PAIN_KEYWORDS.items():
            for keyword in keywords[:2]:  # æ¯ä¸ªäº§å“é€‰2ä¸ªå…³é”®è¯
                search_queries.append(f'"{product}" {keyword}')
        
        # éšæœºé€‰æ‹©5ä¸ªæŸ¥è¯¢
        selected_queries = random.sample(search_queries, min(5, len(search_queries)))
        print(f"  ğŸ¯ ä»Šæ—¥æœç´¢è¯: {selected_queries}")
        
        for query in selected_queries:
            try:
                print(f"  ğŸ” æœç´¢: {query}")
                tweets = await client.search_tweet(query, product='Latest', count=3)
                
                if not tweets:
                    continue
                
                for tweet in tweets:
                    text = tweet.text.replace('\n', ' ')
                    user = tweet.user.name if tweet.user else "Unknown"
                    
                    # æå–äº§å“å
                    product = None
                    for prod in PAIN_KEYWORDS.keys():
                        if prod.lower() in query.lower():
                            product = prod
                            break
                    
                    if product and save_pain("Twitter", user, text, product):
                        count += 1
                
                await asyncio.sleep(1)
            except Exception as e:
                print(f"     âš ï¸ æœç´¢å‡ºé”™: {e}")
                continue
                
    except Exception as e:
        print(f"âŒ Twitter æ‰«æå¤±è´¥: {e}")
    
    return count

def scan_hacker_news():
    """æ‰«æHacker News"""
    print("\nğŸ“° [2/3] æ­£åœ¨æ‰«æ Hacker News...")
    count = 0
    
    try:
        # è·å–çƒ­é—¨æ•…äº‹
        resp = requests.get(
            'https://hacker-news.firebaseio.com/v0/topstories.json',
            timeout=10
        )
        top_ids = resp.json()[:10]
        
        for item_id in top_ids:
            try:
                item = requests.get(
                    f'https://hacker-news.firebaseio.com/v0/item/{item_id}.json',
                    timeout=10
                ).json()
                
                if item and item.get('score', 0) >= 100:
                    title = item.get('title', '')
                    text = item.get('text', '')
                    
                    # æ£€æŸ¥æ˜¯å¦åŒ…å«ç—›ç‚¹å…³é”®è¯
                    for product, keywords in PAIN_KEYWORDS.items():
                        for keyword in keywords:
                            if keyword.lower() in (title + text).lower():
                                content = f"Title: {title} | Text: {text[:100]}"
                                if save_pain("HackerNews", "Tech", content, product):
                                    count += 1
                                break
                
                time.sleep(0.5)
            except:
                pass
                
    except Exception as e:
        print(f"âŒ HN æ‰«æå¤±è´¥: {e}")
    
    return count

def analyze_opportunities(raw_data):
    """ç”¨AIåˆ†æå¸‚åœºæœºä¼š"""
    print("\nğŸ§  [3/3] AI æ­£åœ¨åˆ†æå¸‚åœºæœºä¼š...")
    
    prompt = f"""
# Role: Market Opportunity Analyst
# Task: ä»ç”¨æˆ·åæ§½ä¸­æå–å¯å•†ä¸šåŒ–çš„å¸‚åœºæœºä¼š

## åˆ†ææ¡†æ¶
1. **ç—›ç‚¹åˆ†ç±»**
   - åŠŸèƒ½ç¼ºé™· (Feature Gap)
   - æ€§èƒ½é—®é¢˜ (Performance Issue)
   - æˆæœ¬é—®é¢˜ (Cost Issue)
   - æ˜“ç”¨æ€§ (UX Issue)

2. **æœºä¼šè¯„ä¼°**
   - å—ä¼—è§„æ¨¡ (Market Size)
   - è§£å†³éš¾åº¦ (Technical Difficulty)
   - å•†ä¸šå¯è¡Œæ€§ (Business Viability)

3. **äº§å“å»ºè®®**
   - å¾®äº§å“å½¢æ€ (Micro-SaaS)
   - æ ¸å¿ƒåŠŸèƒ½ (MVP)
   - å®šä»·ç­–ç•¥ (Pricing)

## ç”¨æˆ·æ•°æ®
{raw_data}

## è¾“å‡ºæ ¼å¼
### ğŸ¯ Top 3 å•†ä¸šæœºä¼š

1. [æœºä¼šåç§°]
   - ç—›ç‚¹: ...
   - å¸‚åœºè§„æ¨¡: ...
   - è§£å†³æ–¹æ¡ˆ: ...
   - å»ºè®®è¡ŒåŠ¨: ...

2. [æœºä¼šåç§°]
   ...

3. [æœºä¼šåç§°]
   ...

### ğŸ“Š æ•°æ®ç»Ÿè®¡
- æ€»ç—›ç‚¹æ•°: ...
- ä¸»è¦äº§å“: ...
- æœ€çƒ­è¯é¢˜: ...
"""
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = gemini_client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt
            )
            return response.text
        except Exception as e:
            print(f"âš ï¸ åˆ†æå°è¯• {attempt+1}/{max_retries} å¤±è´¥: {e}")
            if attempt < max_retries - 1:
                time.sleep(5)
            else:
                return "âŒ AIåˆ†æå¤±è´¥ï¼Œè¯·æ£€æŸ¥APIå¯†é’¥"
    
    return "âŒ AIåˆ†æå¤±è´¥"

def deliver_report(content):
    """äº¤ä»˜æŠ¥å‘Š"""
    if content.startswith("âŒ"):
        print(f"\nğŸš« {content}")
        return
    
    today = datetime.date.today().strftime("%Y-%m-%d")
    filename = f"Market_Opportunities_{today}.docx"
    
    # ä¿å­˜Wordæ–‡æ¡£
    try:
        doc = Document()
        doc.add_heading(f'ğŸ¯ å¸‚åœºæœºä¼šåˆ†ææŠ¥å‘Š - {today}', 0)
        doc.add_paragraph(f"ç”Ÿæˆæ—¶é—´: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"æ•è·ç—›ç‚¹æ•°: {len(current_session_pains)}")
        doc.add_paragraph("=" * 50)
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=3)
            else:
                doc.add_paragraph(line)
        
        doc.save(filename)
        print(f"\nğŸ’¾ âœ… æŠ¥å‘Šå·²ç”Ÿæˆ: {filename}")
    except Exception as e:
        print(f"âŒ Wordç”Ÿæˆå¤±è´¥: {e}")
    
    # æ¨é€å¾®ä¿¡
    if PUSHPLUS_TOKEN and PUSHPLUS_TOKEN != 'å¡«å…¥è‡ªå·±çš„TOKEN':
        try:
            print("ğŸ“¨ æ­£åœ¨æ¨é€åˆ°å¾®ä¿¡...")
            wechat_body = f"# ğŸ¯ å¸‚åœºæœºä¼šåˆ†æ ({today})\n\n{content}"
            
            requests.post(
                'http://www.pushplus.plus/send',
                json={
                    "token": PUSHPLUS_TOKEN,
                    "title": f"ã€å¸‚åœºæœºä¼šã€‘{today}",
                    "content": wechat_body,
                    "template": "markdown"
                },
                timeout=10
            )
            print("ğŸ“¨ âœ… å¾®ä¿¡æ¨é€å®Œæˆ")
        except Exception as e:
            print(f"âš ï¸ æ¨é€å¤±è´¥: {e}")

# ==================== ä¸»ç¨‹åº ====================

async def main():
    print("\n" + "="*60)
    print("ğŸš€ å¼€å§‹ç›‘æ§å¾ªç¯")
    print("="*60)
    
    # æ‰§è¡Œæ‰«æ
    c1 = await scan_twitter()
    c2 = scan_hacker_news()
    
    total = c1 + c2
    print(f"\nğŸ“Š æœ¬æ¬¡æ•è·ç—›ç‚¹æ•°: {total}")
    
    if total > 0:
        # æ ¼å¼åŒ–ç—›ç‚¹æ•°æ®
        raw_pains = "\n".join([
            f"ã€{p['source']}ã€‘({p['product']}) @{p['author']}: {p['content']}"
            for p in current_session_pains
        ])
        
        # AIåˆ†æ
        analysis = analyze_opportunities(raw_pains)
        
        # äº¤ä»˜æŠ¥å‘Š
        deliver_report(analysis)
    else:
        print("ğŸ¤· æœªæ•è·åˆ°æ–°ç—›ç‚¹")
    
    print("\nâœ… ç›‘æ§å¾ªç¯å®Œæˆ")

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\nâš ï¸ ç”¨æˆ·ä¸­æ–­")
    except Exception as e:
        print(f"\nâŒ ç¨‹åºé”™è¯¯: {e}")
        import traceback
        traceback.print_exc()
